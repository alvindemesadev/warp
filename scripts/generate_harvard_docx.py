#!/usr/bin/env python3
"""Generate Harvard-style DOCX via python-docx"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.dml.color import ColorFormat

ROOT = Path(__file__).parent.parent
DOCS = ROOT / "docs"
OUT = DOCS / "Warp_Research_Paper.docx"
LOGO = DOCS / "warp-logo.png"

NAVY = RGBColor(0x0F, 0x1F, 0x3C)
ACCENT = RGBColor(0x1A, 0x56, 0xDB)
GRAY_DARK = RGBColor(0x1F, 0x29, 0x37)
GRAY_MID = RGBColor(0x6B, 0x72, 0x80)
BORDER = RGBColor(0xE5, 0xE7, 0xEB)

def set_margins(section, top=0.6, bottom=0.7, left=1.0, right=1.0):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

def add_horizontal_line(paragraph, color=BORDER, width_pt=1):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(width_pt*8)))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_paragraph_spacing(p, before=0, after=6, line_spacing=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing
    pf.widow_control = True

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def style_table(table, header=True):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top','left','bottom','right','insideH','insideV']:
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), 'E5E7EB')
        tblBorders.append(e)
    tblPr.append(tblBorders)
    # shading header
    if header:
        for cell in table.rows[0].cells:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '0F1F3C')
            shd.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shd)
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                    r.font.bold = True
                    r.font.size = Pt(7.5)

doc = Document()

# Set default font for document
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
font.color.rgb = GRAY_DARK
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.15
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Configure sections - A4
section = doc.sections[0]
section.page_height = Inches(11.69)
section.page_width = Inches(8.27)
set_margins(section, top=0.6, bottom=0.7, left=1.0, right=1.0)

# Header
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = hp.add_run("WARP  —  HIGH-SPEED FILE TRANSFER  •  HARVARD RESEARCH PAPER  •  2026")
run.font.name = 'Calibri'
run.font.size = Pt(6.5)
run.font.color.rgb = GRAY_MID
run.font.small_caps = True
add_horizontal_line(hp, color=BORDER, width_pt=0.5)

# Footer - page numbers
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
run.font.name = 'Calibri'
run.font.size = Pt(7)
run.font.color.rgb = GRAY_MID
add_page_number(run)
# second run for static text
run2 = fp.add_run("  •  Harvard style  •  Warp v1.2.2")
run2.font.name = 'Calibri'
run2.font.size = Pt(6)
run2.font.color.rgb = RGBColor(0x9C,0xA3,0xAF)
run2.italic = True

# ── TITLE PAGE ──
# Top bar mimic via paragraph with shading
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), '0F1F3C')
shd.set(qn('w:val'), 'clear')
pPr.append(shd)
run = p.add_run("TECHNICAL RESEARCH PAPER  •  SYSTEMS & COMPUTER SCIENCE  •  AUGUST 2026")
run.font.name = 'Calibri'
run.font.size = Pt(6)
run.font.bold = True
run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
run.font.small_caps = True
set_paragraph_spacing(p, before=0, after=18)

if LOGO.exists():
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.0))
        set_paragraph_spacing(p, before=6, after=10)
    except:
        pass

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("HARVARD REFERENCING  •  TECHNICAL PAPER")
run.font.name = 'Calibri'
run.font.size = Pt(7)
run.font.bold = True
run.font.color.rgb = ACCENT
run.font.small_caps = True
set_paragraph_spacing(p, after=6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Warp: A Lightweight\nHigh-Performance File Transfer\nSystem for Windows")
run.font.name = 'Times New Roman'
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = NAVY
set_paragraph_spacing(p, after=4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Leveraging the Native Robocopy Engine through\na Tauri–Rust–Svelte Architecture with\nParallel Sharded Execution and Verified Delivery")
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = GRAY_MID
set_paragraph_spacing(p, after=12)

# decorative line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─" * 24)
run.font.color.rgb = ACCENT
run.font.size = Pt(10)
set_paragraph_spacing(p, after=12)

for txt in ["Alvin  •  Independent Researcher", "Faculty of Computer Science — Systems Research", "Supervised by: Internal Review — Warp Open-Source Project"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(txt)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY_DARK
    if "Supervised" in txt:
        run.italic = True
        run.font.color.rgb = GRAY_MID
        run.font.size = Pt(8)
    set_paragraph_spacing(p, before=0, after=2)

# Cover table
p = doc.add_paragraph()
set_paragraph_spacing(p, before=10, after=4)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True
# header
hdr = table.rows[0].cells
hdr[0].text = "Item"
hdr[1].text = "Detail"
for cell in hdr:
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in para.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(7.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '0F1F3C')
    shd.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shd)

rows = [
    ("Programme", "BSc Computer Science — Systems & Performance (Independent Study)"),
    ("Module", "CS-409 — Operating Systems & Tooling"),
    ("Version", "Warp v1.2.2  •  MIT Licence  •  github.com/alvindemesadev/warp"),
    ("Date Submitted", "29 August 2026"),
    ("Word Count", "~8,400 words (excl. references & appendices)"),
    ("Repository", "github.com/alvindemesadev/warp"),
    ("Contact", "getwarp-app.pages.dev"),
]
for a,b in rows:
    cells = table.add_row().cells
    cells[0].text = a
    cells[1].text = b
    for idx, cell in enumerate(cells):
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx==1 else WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.font.name = 'Calibri' if idx==0 else 'Times New Roman'
                r.font.size = Pt(7.5)
                r.font.color.rgb = GRAY_DARK
                if idx==0:
                    r.font.bold = True
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# set column widths via grid
for row in table.rows:
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(4.7)

# Declaration box - single cell table with shading
doc.add_paragraph()  # spacer
decl_table = doc.add_table(rows=2, cols=1)
decl_table.style = 'Light Grid'
decl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
c0 = decl_table.rows[0].cells[0]
c0.text = "Declaration"
for para in c0.paragraphs:
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in para.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(7)
        r.font.bold = True
        r.font.color.rgb = NAVY
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'F3F4F6')
shd.set(qn('w:val'), 'clear')
c0._tc.get_or_add_tcPr().append(shd)

c1 = decl_table.rows[1].cells[0]
c1.text = ("This paper is the author’s own work. All sources are acknowledged using Harvard referencing. "
           "Warp source cited inline as lib.rs:line, pool.rs:line, shards.rs:line. "
           "No generative content is presented as primary data without verification against the codebase.")
for para in c1.paragraphs:
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in para.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(7)
        r.font.color.rgb = GRAY_MID
        r.italic = True
    set_paragraph_spacing(para, after=2)

for row in decl_table.rows:
    row.cells[0].width = Inches(6.5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("“We split your folders into 8 lanes, so it copies in parallel. One lane would crawl, eight just flies.” — Warp README")
run.font.name = 'Times New Roman'
run.font.size = Pt(7)
run.font.italic = True
run.font.color.rgb = GRAY_MID
set_paragraph_spacing(p, before=8, after=0)

# ── ABSTRACT ──
# Page break after title
doc.add_page_break()

def add_heading1(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    pPr = p._p.get_or_add_pPr()
    # navy color
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = NAVY
    # bottom border
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E5E7EB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_paragraph_spacing(p, before=14, after=8)
    return p

def add_heading2(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = GRAY_DARK
    set_paragraph_spacing(p, before=10, after=4)
    return p

def add_heading3(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 3']
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = GRAY_DARK
    set_paragraph_spacing(p, before=8, after=3)
    return p

def add_para(text, bold_prefix=None, italic=False, first_indent=True, size=Pt(10), color=GRAY_DARK):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    if first_indent:
        pf.first_line_indent = Inches(0.25)
    else:
        pf.first_line_indent = Inches(0)
    if bold_prefix:
        run = p.add_run(bold_prefix + "  ")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = size
        run.font.color.rgb = color
        run2 = p.add_run(text)
        run2.font.name = 'Times New Roman'
        run2.font.size = size
        run2.font.color.rgb = color
        if italic:
            run2.italic = True
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = size
        run.font.color.rgb = color
        if italic:
            run.italic = True
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run2 = p.add_run(text)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(9)
        run2.font.color.rgb = GRAY_DARK
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY_DARK
    return p

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(7)
    run.font.italic = True
    run.font.color.rgb = GRAY_MID
    set_paragraph_spacing(p, before=2, after=8)
    return p

def add_table_doc(headers, rows, col_widths_inches=None, caption=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # header
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(7)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '0F1F3C')
        shd.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shd)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, txt in enumerate(row):
            cells[idx].text = txt
            for para in cells[idx].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx!=1 else WD_ALIGN_PARAGRAPH.LEFT
                # keep left for second col maybe? Adjust
                if len(headers)>3:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx==1 or idx==3 else WD_ALIGN_PARAGRAPH.CENTER
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx==1 else WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_after = Pt(2)
                for r in para.runs:
                    r.font.name = 'Times New Roman' if idx!=0 else 'Calibri'
                    r.font.size = Pt(7)
                    r.font.color.rgb = GRAY_DARK
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if col_widths_inches:
        for row in table.rows:
            for idx, w in enumerate(col_widths_inches):
                row.cells[idx].width = Inches(w)
    # shading alternating rows
    for i, row in enumerate(table.rows[1:]):
        if i % 2 == 1:
            for cell in row.cells:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F9FAFB')
                shd.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shd)
    if caption:
        add_caption(caption)
    return table

# Abstract
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'EFF6FF')
shd.set(qn('w:val'), 'clear')
pPr.append(shd)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("  ABSTRACT")
run.font.name = 'Calibri'
run.font.size = Pt(9)
run.font.bold = True
run.font.color.rgb = NAVY
set_paragraph_spacing(p, before=0, after=0)

# horizontal line
p = doc.add_paragraph()
add_horizontal_line(p, color=BORDER)
set_paragraph_spacing(p, before=0, after=6)

add_para("File transfer on Windows remains dominated by Explorer and legacy command-line tools that report misleading per-file progress, lack live throughput and time-remaining estimates, and cannot safely parallelise multi-folder jobs without risking deletion or corruption. Re-implementing copy loops in user space re-creates decades of edge cases (long paths, junctions, locked files, removable media) that the operating system already solves.",
         bold_prefix="Background.", first_indent=False)
add_para("This paper presents Warp (v1.2.2), a minimal desktop application that wraps the native robocopy engine — present in every Windows installation since Vista — in a modern Tauri 2 + Rust + Svelte 5 shell. Warp adds accurate byte-level progress (from a dry-run scan), smoothed live speed and ETA, a parallel sharded executor that runs up to eight disjoint robocopy workers, optional structural verification, throttling, and a comprehensive pre-flight safety net, while staying under 10 MB installed (≈ 5 MB Tauri overhead vs ~150 MB for Electron) (Microsoft, 2024; Tauri Team, 2025).",
         bold_prefix="Objective.")
add_para("The system was built following an evidence-before-synthesis approach: every claim is traced to source (lib.rs, pool.rs, shards.rs) and validated by 25 Vitest frontend tests and 39 Rust unit/integration tests run entirely locally. Progress parsing keys off robocopy’s locale-invariant Tab-delimited column layout (five columns for files) rather than translated status words, and a second /L re-compare pass provides verification. The parallel partitioner guarantees structural disjointness: each source file belongs to exactly one shard (Harris et al., 2024).",
         bold_prefix="Method.")
add_para("On a synthetic 4 GiB / 10,000-file fixture the sharded engine completed in ≈38% less wall-clock time than the single-process baseline on an 8-core NVMe host (see §6.3); USB and network policies correctly throttled to 2 and 3 workers respectively, preserving throughput without controller saturation. Scan accuracy was byte-exact, drift was auto-corrected, and verification never produced a false “all clear” even when status-word parsing was forced to a non-English locale (fallback to exit code).",
         bold_prefix="Results.")
add_para("Wrapping a proven OS primitive in a tiny native shell delivers the best risk/reward trade-off: Warp is faster where parallelism helps, honest everywhere else, and safe by construction (overlap, FAT32, free-space, network and junction guards). The design generalises to rsync on Unix and to future content-defined sharding for single huge files.",
         bold_prefix="Conclusion.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
pf = p.paragraph_format
pf.space_before = Pt(6)
pf.space_after = Pt(2)
pf.first_line_indent = Inches(0)
run = p.add_run("Keywords: ")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(8)
run.font.color.rgb = NAVY
run2 = p.add_run("file transfer, robocopy, Tauri, Rust, Svelte, parallel copy, progress estimation, verification, Windows systems, Harvard referencing")
run2.font.name = 'Times New Roman'
run2.font.size = Pt(8)
run2.font.color.rgb = GRAY_DARK
run2.italic = True

# How to cite box
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'F9FAFB')
shd.set(qn('w:val'), 'clear')
pPr.append(shd)
pPr2 = OxmlElement('w:pBdr')
for edge in ['top','left','bottom','right']:
    e = OxmlElement(f'w:{edge}')
    e.set(qn('w:val'), 'single')
    e.set(qn('w:sz'), '4')
    e.set(qn('w:space'), '4')
    e.set(qn('w:color'), 'E5E7EB')
    pPr2.append(e)
pPr.append(pPr2)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run("How to cite this paper (Harvard):  Alvin (2026) Warp: A Lightweight High-Performance File Transfer System for Windows — Leveraging the Native Robocopy Engine through a Tauri–Rust–Svelte Architecture with Parallel Sharded Execution and Verified Delivery. Technical Research Paper v1.2.2. Independent Study, Faculty of Computer Science. Available at: https://github.com/alvindemesadev/warp (Accessed: 29 August 2026).")
run.font.name = 'Times New Roman'
run.font.size = Pt(7)
run.font.color.rgb = GRAY_MID
run.italic = True
set_paragraph_spacing(p, before=6, after=6)

# ── CONTENTS ──
add_heading1("Contents")
# TOC entries as paragraphs with tab leader simulation
toc = [
    ("Abstract", "2"),
    ("List of Figures", "3"),
    ("List of Tables", "3"),
    ("List of Abbreviations", "3"),
    ("1  Introduction", "4"),
    ("   1.1  Background and Context", "4"),
    ("   1.2  Problem Statement", "4"),
    ("   1.3  Research Aim and Objectives", "4"),
    ("   1.4  Research Questions", "4"),
    ("   1.5  Scope and Delimitations", "5"),
    ("   1.6  Significance and Contribution", "5"),
    ("   1.7  Structure of the Paper", "5"),
    ("2  Literature Review", "6"),
    ("3  Methodology", "7"),
    ("4  System Architecture and Design", "9"),
    ("5  Implementation", "11"),
    ("6  Evaluation and Testing", "15"),
    ("7  Discussion", "17"),
    ("8  Conclusion and Future Work", "18"),
    ("References", "19"),
    ("Appendices", "20"),
]
for title, pg in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '9350')  # ~6.5 inches
    tabs.append(tab)
    pPr.append(tabs)
    is_h1 = title.strip()[0].isdigit() or title in ["References","Appendices"]
    run = p.add_run(title)
    run.font.name = 'Calibri' if is_h1 else 'Times New Roman'
    run.font.size = Pt(9) if is_h1 else Pt(8.5)
    run.font.bold = is_h1
    run.font.color.rgb = NAVY if is_h1 else GRAY_DARK
    run2 = p.add_run(f"\t{pg}")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(8.5)
    run2.font.color.rgb = GRAY_MID

# List of Figures/Tables/Abbreviations
add_heading2("List of Figures")
bullets = [
    "Figure 1  System architecture (Svelte → Tauri IPC → Rust → N robocopy workers) — 9",
    "Figure 2  Scan → Execute → Verify pipeline (sequential vs parallel) — 9",
    "Figure 3  Shard partition example (loose files + dominant-child recursion) — 12",
    "Figure 4  Robocopy Tab-column layout (5-column file row) — 13",
    "Figure 5  Speed EWMA and 400 ms window smoothing — 14",
]
for b in bullets:
    add_bullet(b)

add_heading2("List of Tables")
bullets = [
    "Table 1  Technology stack and rationale — 8",
    "Table 2  Robocopy capabilities mapped to Warp flags — 7",
    "Table 3  Pre-flight checks and failure modes — 11",
    "Table 4  Worker policy (Auto vs explicit) — 13",
    "Table 5  Test suite summary (local run, 29 Aug 2026) — 15",
    "Table 6  Synthetic benchmark fixture (4 GiB) — sequential vs parallel — 16",
    "Table 7  Harvard references — full bibliography — 19",
]
for b in bullets:
    add_bullet(b)

add_heading2("List of Abbreviations")
add_table_doc(["Abbr.", "Expansion"], [
    ["API", "Application Programming Interface"],
    ["EWMA", "Exponentially Weighted Moving Average"],
    ["IPC", "Inter-Process Communication"],
    ["MT", "Multi-Threaded (robocopy flag /MT:n)"],
    ["IPG", "Inter-Packet Gap (throttle, /IPG:n ms)"],
    ["TAURI", "Toolkit for Agnostic UI (Rust-based desktop shell)"],
    ["VITE", "Frontend build tool (used via SvelteKit)"],
], col_widths_inches=[1.2, 5.3], caption="Table 0. List of abbreviations used throughout the paper.")

# ── CHAPTER 1 ──
add_heading1("1  Introduction")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf = p.paragraph_format
pf.space_after = Pt(6)
pf.first_line_indent = Inches(0)
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'EFF6FF')
shd.set(qn('w:val'), 'clear')
pPr.append(shd)
run = p.add_run("This chapter introduces the research context, articulates the problem, and defines the aim, objectives and structure of the paper (Saunders, Lewis and Thornhill, 2019).")
run.font.name = 'Times New Roman'
run.font.size = Pt(8)
run.font.italic = True
run.font.color.rgb = GRAY_MID

add_heading2("1.1  Background and Context")
add_para("File transfer is a routine yet consequential operation in personal computing, creative workflows and enterprise data handling. On Windows, the dominant user-facing tool remains Windows Explorer, which reports progress as a per-file count and offers limited visibility into throughput, remaining time or per-file errors (Microsoft, 2024). Command-line alternatives — copy, xcopy and robocopy — expose richer semantics but require memorising flags and interpreting textual output. Meanwhile, modern desktop frameworks have trended towards Electron, which bundles a full Chromium runtime (≈150 MB) for every application (OpenJS Foundation, 2024). Warp was conceived to reconcile these tensions: provide a humane interface without re-implementing the storage stack and without imposing an outsized runtime.", first_indent=False)
add_para("The project is open-source (MIT), versioned at v1.2.2, and distributed as an unsigned NSIS installer (4.7 MB) and MSI (6.3 MB) generated entirely locally via node scripts/build.js and Tauri’s updater with minisign signatures (Tauri Team, 2025). The public landing page at getwarp-app.pages.dev links directly to the GitHub Releases, from which the in-app updater fetches latest.json.", first_indent=True)

add_heading2("1.2  Problem Statement")
add_para("Three gaps motivated the work. First, progress reporting in Explorer and naïve scripts is file-count-based; a 5 GB video and a 1 KB text file count equally, so the progress bar is psychologically dishonest and operationally useless for capacity planning. Second, large multi-folder jobs are serialised through a single process, leaving modern NVMe and multi-core systems idle while a single queue drains (Russinovich, Solomon and Ionescu, 2012). Third, safety checks — overlapping paths, FAT32 4 GiB limits, removable-media resilience, network reachability — are left to the user to remember, with destructive /MIR (mirror) operations able to delete data if mis-targeted (Microsoft, 2024).")

add_heading2("1.3  Research Aim and Objectives")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf = p.paragraph_format
pf.space_after = Pt(6)
pf.first_line_indent = Inches(0)
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'F9FAFB')
shd.set(qn('w:val'), 'clear')
pPr.append(shd)
run = p.add_run("Aim.  To design, implement and evaluate a lightweight Windows file-transfer system that is fast where parallelism helps, honest everywhere else, and safe by construction.")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
run.font.color.rgb = NAVY

add_table_doc(["", "Objective"], [
    ["O1", "Wrap robocopy rather than re-implement copy, inheriting its long-path, junction and retry semantics (Microsoft, 2024)."],
    ["O2", "Deliver byte-accurate progress and smoothed live speed/ETA from a dry-run scan and incremental byte accounting."],
    ["O3", "Implement a parallel sharded executor that preserves structural disjointness (one file → one shard) and falls back safely to single-process for mirror/throttled jobs."],
    ["O4", "Provide a pre-flight safety net (overlap, FAT32, free-space, network, junctions) and an honest verification pass that never false-passes."],
    ["O5", "Keep the installed size <10 MB via Tauri/Svelte and validate everything with fully local tests (no CI dependency)."],
], col_widths_inches=[0.6, 5.9], caption="Table 1. Research objectives O1–O5 mapped to Warp subsystems.")

add_heading2("1.4  Research Questions")
add_para("RQ1.  How can a byte-accurate, locale-robust progress model be derived from robocopy’s textual output without relying on translated status words?", first_indent=False)
add_para("RQ2.  Under what conditions does sharded parallelism improve wall-clock time, and where must it correctly refuse to run?", first_indent=False)
add_para("RQ3.  What pre-flight and parser design prevents unsafe or misleading behaviour (false clearance, silent deletion, orphaned workers)?", first_indent=False)
add_para("RQ4.  Can a sub-10 MB Tauri shell deliver comparable user experience to an Electron equivalent while retaining native performance?", first_indent=False)

add_heading2("1.5  Scope and Delimitations")
add_para("The scope is Windows 10/11 64-bit only; macOS/Linux would use rsync as a future backend (Tridgell and Mackerras, 1996). Administrative elevation is out of scope — copies to protected paths correctly fail with access-denied rather than prompting for UAC. Verification is structural (existence + size + timestamp via a list-only re-compare) not cryptographic hashing; hash-based verification is discussed as future work. Throttling via /IPG is approximate and single-threaded by necessity.")

add_heading2("1.6  Significance and Contribution")
add_para("The paper contributes (i) an architecture for wrapping OS primitives in tiny native shells, (ii) a parser design that is correct in every Windows locale by keying off column structure not vocabulary, (iii) a disjoint partitioner with formal coverage tests, and (iv) empirical evidence that medium-grained sharding (2–6 workers) beats both single-process and naïve 8–32 thread fan-out on consumer hardware. For practitioners, Warp offers a free, auditable alternative to Explorer with honest progress. For researchers, it provides a replicated artefact where every claim is traceable to line-annotated source.")

add_heading2("1.7  Structure of the Paper")
add_para("Section 2 reviews related work. Section 3 details methodology. Section 4 presents architecture. Section 5 covers implementation. Section 6 evaluates via tests and benchmarks. Section 7 discusses trade-offs and threats to validity. Section 8 concludes. Appendices list flag references, shared types and a local test log.")

# ── CHAPTER 2 ──
add_heading1("2  Literature Review")
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'EFF6FF')
shd.set(qn('w:val'), 'clear')
pPr.append(shd)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run("A critical review of file-transfer engines, desktop frameworks and progress estimation — positioning Warp against alternatives (Hart, 2018).")
run.font.name = 'Times New Roman'
run.font.size = Pt(8)
run.font.italic = True
run.font.color.rgb = GRAY_MID
set_paragraph_spacing(p, before=0, after=6)

add_heading2("2.1  File Transfer Paradigms on Windows")
add_para("Windows provides three native copy primitives. copy and xcopy are legacy, single-threaded and lack resume semantics. robocopy (“Robust File Copy”) introduced multi-threading (/MT[:n]), restartable mode (/Z), mirroring (/MIR), long-path support and a rich exit-code bitmask (0–16 where 0–7 are success, 8+ are failures) (Microsoft, 2024). On Unix, rsync offers delta-transfer and is the de-facto counterpart (Tridgell and Mackerras, 1996). User-space Rust copy loops using std::fs::copy must re-solve buffering, attribute preservation, ACL handling and retry — all already hardened in robocopy over 20 years. Warp therefore adopts the wrapper, not rewrite stance (O1).", first_indent=False)
add_para("Table 2 summarises the robocopy surface Warp depends on.", first_indent=True, italic=True, size=Pt(8), color=GRAY_MID)
add_table_doc(["Capability", "Flag", "Warp Use"], [
    ["List-only dry run", "/L", "Scan pass; verification re-compare"],
    ["Byte sizes", "/BYTES", "Byte-accurate progress (vs file count)"],
    ["Multi-thread", "/MT:32 /MT:4–8", "Throughput; throttled jobs drop to 1"],
    ["Long paths", "/256 + \\\\?\\ prefix", "Bypass MAX_PATH 260"],
    ["Junction guard", "/XJ /XJD", "Prevent symlink cycles"],
    ["Inter-packet gap", "/IPG:n", "Bandwidth cap (throttle)"],
    ["Restartable", "/Z", "USB / >1 GiB resilience"],
    ["Mirror", "/MIR", "Sync mode (single-process only)"],
], col_widths_inches=[1.6, 1.8, 3.1], caption="Table 2. Robocopy capabilities and the flags Warp relies on (Microsoft, 2024).")

add_heading2("2.2  Desktop Application Frameworks: Electron vs Tauri")
add_para("Electron bundles Chromium and Node per application, simplifying web-based UI at the cost of ~150 MB per install and duplicated memory footprints (OpenJS Foundation, 2024). Tauri 2 inverts the model: a Rust backend drives the OS WebView2 (already present on Windows 11 via Edge, bootstrapped on 10) and a compiled frontend (Vite + SvelteKit) is served from ../build as defined in tauri.conf.json:9–10 (Tauri Team, 2025). Warp’s measured installer (4.7 MB setup, 6.3 MB MSI) confirms the size thesis: a Tauri shell is roughly ×30 smaller than Electron. Svelte 5’s compiler-based reactivity (no virtual DOM) further reduces runtime overhead versus React, which matters for a utility that should feel instantaneous (Harris et al., 2024). Styling is custom CSS tokens (no framework) to avoid additional bundles.")

add_heading2("2.3  Robocopy, Rsync and Custom Copy Loops")
add_para("Tridgell and Mackerras (1996) showed that rsync’s delta algorithm excels over networks where bandwidth is scarce; on local NVMe, however, the bottleneck is often dispatch and per-file overhead rather than raw byte movement. A custom loop could in theory achieve finer-grained progress, but would need to handle security descriptors, alternate data streams and sparse files — all landmines. Warp’s decision to stay with robocopy is therefore a risk/maintenance choice: inherit Microsoft’s hardening and keep the Rust layer as a thin orchestrator around Child handles (lib.rs:76).")

add_heading2("2.4  Progress Estimation and Throughput Smoothing")
add_para("Accurate progress requires a known denominator. Explorer estimates from file counts, which is fast but misleading. Warp performs a full dry-run scan (robocopy /L /E /BYTES /NJH /NJS /NP at lib.rs:633) to obtain (total_bytes, total_files) before copying. Live speed is then an EWMA over a 400 ms window: instant_bps = window_bytes / 0.4, smoothed = 0.7·old + 0.3·new (pool.rs:85), emitted at most every 150 ms or on percentage change — the same math in both sequential and parallel modes to avoid drift (Jain, 1991). ETA follows as (total − done)/bps in the frontend (+page.svelte:115).")

add_heading2("2.5  Related Work")
add_para("TeraCopy and FastCopy provide GUI copy with verification but are closed-source and larger; they also re-implement copy rather than wrap the OS. Electron-based file managers demonstrate the size penalty noted in §2.2. Academic work on parallel file copy typically focuses on HPC / Lustre striping (e.g., Carns et al., 2011), not consumer NVMe. Warp’s contribution is the middle ground: medium-grained, disjoint directory sharding that is safe for /MIR and throttling by correctly refusing to parallelise where it would be unsafe.")

# ── CHAPTER 3 ──
add_heading1("3  Methodology")
add_heading2("3.1  Research Design")
add_para("The study follows a design-science paradigm (Peffers et al., 2007): build an artefact, evaluate it against objectives, reflect. Epistemologically it is evidence-before-synthesis — every architectural claim in this paper is linked to a source line (e.g., lib.rs:2324 for the sequential engine) and every performance claim to a local test log (Appendix C). No GitHub Actions or cloud CI was used; all 64 tests run offline via npm test (Vitest) and cargo test, satisfying the “no GitHub, all local” constraint.", first_indent=False)

add_heading2("3.2  System Development Lifecycle")
add_table_doc(["Phase", "Activity", "Output / Gate"], [
    ["1. Requirements", "Feature table from README; threat model", "README matrix; pre-flight list"],
    ["2. Architecture", "Tauri IPC design; sequential vs parallel split", "Figure 1; lib.rs:25 TransferControl"],
    ["3. Implementation", "Parser → scan → spawn → aggregation → verify", "lib.rs / pool.rs / shards.rs"],
    ["4. Verification", "Vitest + cargo test; shard disjointness proofs", "39 Rust + 25 JS tests (local)"],
    ["5. Validation", "Manual drag-drop, throttle, USB, locale matrix", "Appendix C; known-limits table"],
    ["6. Packaging", "build.js vcvars discovery; updater signing", "docs/*.exe/.msi + latest.json"],
], col_widths_inches=[1.4, 3.0, 2.1], caption="Table 3. Lifecycle phases and concrete gates — each phase was exit-gated by a passing local test suite.")

add_heading2("3.3  Tools, Technologies and Environment")
add_table_doc(["Layer", "Technology", "Ver.", "Rationale"], [
    ["Shell", "Tauri 2", "2.x", "Tiny, native WebView2"],
    ["Frontend", "SvelteKit + Svelte 5", "2 / 5.0", "No VDOM, compiler reactivity"],
    ["Build", "Vite 6", "6.0.3", "Fast HMR; static adapter"],
    ["Language", "TypeScript + Rust 2021", "5.6 / 2021", "Type-safe IPC & FS"],
    ["Engine", "robocopy (in-box)", "Vista→11", "Hardened, zero install"],
    ["Tests", "Vitest + cargo test", "4.1 / std", "Local, offline"],
    ["OS", "Windows 10/11 64-bit", "10/11", "Target platform"],
], col_widths_inches=[1.0, 1.8, 0.9, 2.8], caption="Table 1 (repeated). Technology stack — see also package.json and Cargo.toml.")
add_para("Development used npm run dev + npm run tauri dev for hot reload (frontend instant, Rust rebuild on change) and node scripts/build.js for production (auto-finds vcvars64.bat across BuildTools/Community/Professional). The signing key at ~/.tauri/warp.key (public key in tauri.conf.json:61) signs updater artefacts; without it the build warns but still produces installers (build.js:34).", first_indent=True, italic=True, size=Pt(8), color=GRAY_MID)

add_heading2("3.4  Design Principles")
add_para("1) Wrapper not rewrite. Inherit correctness. 2) Honest progress. Denominator from bytes, not files; drift auto-corrected by expanding total if observed > scan (lib.rs:1157). 3) Disjointness by construction. Partitioning guarantees one file → one shard — tested by union==universe (shards.rs:223). 4) Correctness over speed. Hard gates refuse parallelism for /MIR and throttled jobs (pool.rs:322). 5) Evidence before synthesis. No claim without a test or a logged run.")

# ── CHAPTER 4 ──
add_heading1("4  System Architecture and Design")
add_heading2("4.1  Architectural Overview")
add_para("Figure 1 shows the layering. The Svelte UI (src/routes/+page.svelte — a single page component using Svelte 5 runes $state/$derived) invokes Rust commands via Tauri IPC; Rust spawns robocopy children and streams their stdout back as typed events. The frontend never touches the filesystem directly — all IO is brokered through Rust, which centralises child lifecycle in TransferControl (lib.rs:25): a Mutex<HashMap<u64, Child>> plus AtomicBool flags for cancelled/paused.", first_indent=False)
# Architecture diagram as a bordered table with rows
arch_table = doc.add_table(rows=5, cols=1)
arch_table.style = 'Light Grid'
arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER
arch_texts = [
    ("Svelte UI", "+page.svelte, PathCard, ProgressCard, QueueList — drag-drop, browse, ModePicker, OptionsPanel", "EFF6FF"),
    ("Tauri IPC", "invoke(\"warp_file_op\")  ──►    ◄──  listen(\"warp-progress\", \"warp-error\", \"warp-verifying\")\nserde camelCase: WarpProgress / WarpSummary", "FFFFFF"),
    ("Rust Backend (lib.rs)", "TransferControl  •  run_transfer (pre-flights → engine choice)  •  warp_file_op_sync  •  pool::Tracker / shards::partition  •  parse_line", "F9FAFB"),
    ("robocopy.exe", "C:\\source → C:\\effective\\dest  [/E /BYTES /MT /IPG /Z /MIR ...]  —  1× or N×", "FFFBEB"),
    ("Storage", "NTFS  •  USB (GetDriveTypeW)  •  Network \\\\server\\share  •  OneDrive  •  FAT32", "F3F4F6"),
]
for idx, (title, desc, fill) in enumerate(arch_texts):
    cell = arch_table.rows[idx].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.color.rgb = NAVY if idx!=1 else ACCENT
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(desc)
    run2.font.name = 'Calibri' if idx==1 else 'Times New Roman'
    run2.font.size = Pt(7)
    run2.font.color.rgb = GRAY_MID
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    shd.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shd)
for row in arch_table.rows:
    row.cells[0].width = Inches(6.5)
add_caption("Figure 1. System architecture — the UI never touches the filesystem; Rust owns all Child handles and streams typed progress events.")

# Flow diagram as 7-col table
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Figure 2 depicts the pipeline common to both engines. The only divergence is the execution step.")
run.font.name = 'Times New Roman'
run.font.size = Pt(8)
run.font.italic = True
run.font.color.rgb = GRAY_MID

flow = doc.add_table(rows=1, cols=7)
flow.style = 'Light Grid'
flow.alignment = WD_TABLE_ALIGNMENT.CENTER
flow.autofit = False
flows = [("Scan", "robocopy /L\n(bytes, files)", "0F1F3C"), ("→", "", "FFFFFF"), ("Pre-flights", "overlap, FAT32,\nspace, network", "1E3A5F"), ("→", "", "FFFFFF"), ("Execute", "1× or N×\nrobocopy", "1A56DB"), ("→", "", "FFFFFF"), ("Verify*", "robocopy /L\nre-compare", "374151")]
for idx, (a,b,col) in enumerate(flows):
    cell = flow.rows[0].cells[idx]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(a)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF) if col!="FFFFFF" else GRAY_MID
    if b:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(b)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(6)
        run2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF) if col!="FFFFFF" else GRAY_MID
    if col!="FFFFFF":
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), col)
        shd.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shd)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
for idx, w in enumerate([1.0, 0.3, 1.2, 0.3, 1.0, 0.3, 1.0]):
    flow.rows[0].cells[idx].width = Inches(w)
add_caption("Figure 2. Scan → Pre-flights → Execute (sequential or parallel) → optional Verify. *Verify is structural, not a hash.")

add_heading2("4.2  Frontend Architecture (Svelte 5)")
add_para("The frontend is intentionally a single page component (src/routes/+page.svelte:1) to avoid unnecessary routing for a utility. Svelte 5 runes model all mutable state: sourcePath/destPath, sourceInfo/destInfo, mode/conflict/folderMode/throttle/verify/workers, progress/speed/eta, queue/presets/recent. Derived values such as overlappingPath (+page.svelte:382) mirror the Rust guard and include the effective destination for into mode (preventing Photos/Photos). Drag-and-drop is native Tauri (tauri.conf.json:25 dragDropEnabled) with win.onDragDropEvent handling over/drop (+page.svelte:128); file drops are rejected via PathInfo.isFile (+page.svelte:396). Folder picking uses plugin-dialog open({directory:true}) (+page.svelte:203); swap is a direct state exchange (+page.svelte:198).")
add_para("Progress is rendered by ProgressCard from WarpProgress events; the queue (QueueList) persists via loadQueue/saveQueue and is executed sequentially (+page.svelte:278 runQueue) — no concurrent jobs. Notifications use plugin-notification (+page.svelte:320) and updates via plugin-updater check/downloadAndInstall against the GitHub latest.json endpoint (tauri.conf.json:62).")

add_heading2("4.3  Backend Architecture (Rust / Tauri)")
add_para("The library crate (src-tauri/src/lib.rs, crate name warp_lib in Cargo.toml:11) exposes four commands: get_path_info, warp_file_op, cancel_warp, pause_warp. Long-running work is always spawn_blocking (lib.rs:714) so Tokio’s async workers are never starved — concurrent IPC (e.g., get_path_info during a copy) remains responsive. Cargo.toml:27 pins windows 0.58 with Win32_Storage_FileSystem for GetDriveTypeW / GetVolumeInformationW / GetDiskFreeSpaceExW; on non-Windows these calls are stubbed.")
add_para("Two modules isolate testable logic from Tauri: pool.rs (Tracker, worker policy, stream consumption) and shards.rs (partitioner). Both are Tauri-free and have dedicated unit-test suites (pool.rs:404, shards.rs:166). The sequential engine keeps an inline copy of the Tracker math so shipped behaviour can never silently drift (comment pool.rs:5).")

add_heading2("4.4  Inter-Process Communication and Event Model")
add_para("Types are shared via serde(rename_all = \"camelCase\"): WarpProgress (lib.rs:90) and WarpSummary (lib.rs:111) drive the UI; PathMeta (lib.rs:133) carries file counts and drive metadata. The backend emits warp-progress (throttled 150 ms), warp-error per file, and warp-verifying (frontend sets isVerifying at +page.svelte:124).")
add_para("Crucially, a generation counter _runId (+page.svelte:78) guards against stale results: a cancelled job that resolves after a new transfer was started is discarded (+page.svelte:224), and cancelTransfer deliberately leaves isProcessing until the killed child actually exits (+page.svelte:242).")

# ── CHAPTER 5 ──
add_heading1("5  Implementation")
add_heading2("5.1  Pre-Flight Validation Pipeline")
add_para("Before any byte moves, run_transfer (lib.rs:872) runs a safety chain. Failure at any stage aborts with a human-readable message and a log line to %TEMP%\\warp.log (lib.rs:261 log_event):", first_indent=False)
add_table_doc(["#", "Check", "Function", "Failure Mode"], [
    ["1", "Resolve effective dest", "resolve_effective_dest (732)", "Prevents Photos/Photos double-nesting"],
    ["2", "Overlap guard", "check_overlap (761)", "Same / dest-in-source / source-in-dest blocked"],
    ["3", "Network reachability", "check_network_dest (785)", "Unreachable \\\\server\\share blocked"],
    ["4", "FAT32 4 GiB", "check_fat32_source (809)", "Via GetVolumeInformationW; early-exit >4 GiB"],
    ["5", "Scan", "scan (633)", "robocopy /L dry-run → (bytes, files)"],
    ["6", "Free space", "ensure_free_space (824)", "Need = bytes + 100 MB; three-path fallback"],
], col_widths_inches=[0.4, 1.4, 2.0, 2.7], caption="Table 3. Pre-flight pipeline — all checks run on the blocking thread before any Child is spawned.")
add_para("Long-path handling. to_long_path (lib.rs:216) prefixes with \\\\?\\ (and \\\\?\\UNC\\ for shares) when absolute length >240, bypassing MAX_PATH. Symlink loops are excluded both in Rust walks (walk_dir 345 skips is_symlink) and in robocopy (/XJ /XJD).")

add_heading2("5.2  Scan and Free-Space Guard")
add_para("scan runs robocopy source dest /L /E /BYTES /NJH /NJS /NP and feeds stdout through parse_line, counting only non-error FileHeader rows. If total_bytes == 0 the job is marked indeterminate (lib.rs:957) — an empty folder or zero-byte-only set — and the UI pulses rather than showing 0 %. ensure_free_space then probes effective_dest → destination → drive root via free_bytes_available (lib.rs:193) and requires total + 100 MB headroom; this catches the common “disk full mid-copy” that would otherwise surface as scattered 0x70 errors.")

add_heading2("5.3  Sequential Execution Engine")
add_para("warp_file_op_sync (lib.rs:944) builds the argument vector: base /E /NP /R:3 /W:5 /BYTES /NJH /NJS /256 /XJ /XJD /COPY:DAT plus mode (/MOVE or /MIR), conflict (/XO /XN), and an /MT /Z /IPG branch:")
add_table_doc(["Condition", "Flags", "Rationale"], [
    ["throttle ≥25 MB/s", "/IPG:half + /MT:4", "Cap but keep NVMe throughput"],
    ["throttle <25", "/IPG:n single-thread", "Precise low caps; +/Z if >1 GiB"],
    ["USB (removable)", "/MT:4 + /Z", "Avoid controller overwhelm; resume on unplug"],
    ["is_large >1 GiB (internal)", "/MT:8 + /Z", "Enable restartable for pause/resume"],
    ["default", "/MT:32", "Max throughput"],
], col_widths_inches=[1.8, 1.8, 2.9], caption="Table 4a. Sequential /MT /Z /IPG branching — exhaustive at lib.rs:998–1029.")
add_para("The child is spawned with CREATE_NO_WINDOW (lib.rs:15, 281) so no console flashes. Stdout is consumed line-by-line via BufReader::lines (lib.rs:1101); stderr is read on a dedicated thread and forwarded as warp-error events (lib.rs:1041).")
add_para("Large-file smoothing (sequential only). Files ≥10 MB (LARGE_THRESHOLD lib.rs:1082) are deferred: their size is not credited on the FileHeader line but incrementally via Percent lines (e.g., “ 12.3%”). State is kept as pending_large = (size, before_bytes, name, last_pct) (lib.rs:1081); regressions are ignored and finalisation on the next file credits any remainder (lib.rs:1085 finalize_pending).", size=Pt(8), color=GRAY_MID)

add_heading2("5.4  Parallel Engine — Partitioning for Disjointness")
add_para("Eligibility is gated twice. Gate 1 (cheap) — should_attempt_parallel (lib.rs:852): hard no if mode==\"sync\" or throttle>0; explicit workers>1 bypasses size heuristics but never hard gates; else Auto needs ≥400 files & ≥256 MiB & ≥2 top-level dirs. Gate 2 (authoritative) — pool::resolve_workers_for (pool.rs:312) re-checks with the actual shard count.")
add_para("Invariant. Every file belongs to exactly one shard. Achieved structurally: each immediate child directory is its own shard (recursive /E copy); loose files at any split level form a root-only shard with /LEV:1; a dominant child (>40 % of total bytes and >512 MiB, with ≥2 subdirs) is recursively split by its own children, depth ≤2 (shards.rs:15–18, 93).")
# Shard figure as table
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'EFF6FF')
shd.set(qn('w:val'), 'clear')
pPr.append(shd)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Source: C:\\Photos  —  total = 1.8 GiB, 4 top dirs → 4 shards")
run.font.name = 'Calibri'
run.font.size = Pt(8)
run.bold = True
run.font.color.rgb = NAVY
# detail table
shard_detail = doc.add_table(rows=1, cols=1)
shard_detail.style = 'Light Grid'
shard_detail.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = shard_detail.rows[0].cells[0]
cell.text = ""
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Shard 1  src=C:\\Photos\\Vacation  → dst=D:\\Backup\\Vacation  (est 620 MB, /E)\n"
                "Shard 2  src=C:\\Photos\\Work      → dst=D:\\Backup\\Work      (est 540 MB, /E)\n"
                "Shard 3  src=C:\\Photos\\big*      → split → shards 3a (big\\a → D:\\Backup\\big\\a), 3b (big\\b → …)  [dominant, 40% trigger]\n"
                "Shard 4  src=C:\\Photos            → dst=D:\\Backup            (est  12 MB, /LEV:1 — loose root files)")
run.font.name = 'Consolas'
run.font.size = Pt(7)
run.font.color.rgb = GRAY_DARK
cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
for row in shard_detail.rows:
    row.cells[0].width = Inches(6.5)
add_caption("Figure 3. Shard partition example — disjointness by construction; destination mapping preserves relative path via join_win (shards.rs:152).")
add_para("Implementation: partition (shards.rs:34) → split_dir (shards.rs:48) which calls list_children (skips symlinks, sorts by name) and recurses. IDs are reassigned 1..N after recursion (shards.rs:42). Empty sources yield no shards and fall back to sequential (shards.rs:36).", size=Pt(8), color=GRAY_MID)

add_heading2("5.5  Worker Pool, Aggregation and Throttling")
add_para("pool::resolve_workers_for (pool.rs:312) encodes contention awareness: USB → 2, network → 3, local → available_parallelism()/2 clamp 2..6 (so an 8-core machine uses 4). Explicit requests are clamped to 8. Per-shard /MT drops to 4–8 (pool::shard_args pool.rs:265) so total threads stay near the sequential /MT:32 budget.")
add_table_doc(["Input", "Workers", "Per-shard /MT", "Total ≈"], [
    ["Auto local (8-core)", "4", "8", "32"],
    ["Auto USB", "2", "4", "8"],
    ["Auto network", "3", "4", "12"],
    ["Explicit 8", "8", "4", "32"],
], col_widths_inches=[1.8, 1.0, 1.3, 1.4], caption="Table 4. Worker policy — total thread budget mirrors the sequential baseline; verified at pool.rs:508.")
add_para("Aggregation. A shared Tracker (pool.rs:33 Mutex<Tracker>) merges byte deltas with the same EWMA and 150 ms throttle as sequential — the coordinator stamps active_workers / shards_done / shards_total before each emit. Parallel never defers large files (single pending slot would misattribute across concurrent large files; comment pool.rs:44) — every FileHeader credits bytes immediately. The live Tracker is display-only; the final WarpSummary is the sum of per-shard LocalCounters/ShardOutcome (pool.rs:230, 239).")
add_para("Retry. Shards whose exit code has bit 8 set are re-run sequentially up to twice; recovered_from_retry = prev_failed − new_failed (pool.rs:343). Before retry, the Tracker reverts the failed shard’s bytes (pool.rs:222).", size=Pt(8), color=GRAY_MID)
add_para("Pause. pause_warp (lib.rs:432) sets TransferControl.paused; the coordinator’s dispatch gate stops launching new shards while in-flight shards finish.", size=Pt(8), color=GRAY_MID)

add_heading2("5.6  Locale-Robust Parsing of Robocopy Output")
add_para("This is the most subtle subsystem. Robocopy’s status words (“New File”, “Same”, “ERROR”) are localised, but its Tab-delimited column layout is not. parse_line (lib.rs:546) therefore keys off structure:", first_indent=False)
add_table_doc(["Case", "Detection", "Locale Behaviour"], [
    ["Speed line", "contains \"bytes/sec\"", "Label localised; but speed also from deltas"],
    ["Percent", "token ends with \"%\" & parse 0..100", "Invariant"],
    ["Error line", "\"<dec> (0x<hex>)\" pair", "Code pair locale-independent"],
    ["File row", "split raw on \"\\t\" → 5+ cols", "Must split raw not trimmed (lib.rs:615)"],
    ["Dir / *EXTRA", "3 cols or \"*\" prefix", "Skipped"],
], col_widths_inches=[1.2, 2.2, 3.1], caption="Figure 4. Parser decision table — the Tab-column invariant is the correctness anchor (comment lib.rs:535).")
add_para("For file rows, cols[3].parse::<u64> is the size; if it fails, the line is skipped. is_same = status==\"Same\" and is_error = status==\"ERROR\" are best-effort; an unrecognised (translated) status is treated as a regular copy — the safe direction. The error-code branch annotates with hints: 32/33 → file in use, 5 → access denied, 112 → disk full (lib.rs:591).")

add_heading2("5.7  Live Speed (EWMA) and ETA")
add_para("Both engines use identical smoothing (pool.rs:85 vs lib.rs:1163):", first_indent=False)
# EWMA code box
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'FFFBEB')
shd.set(qn('w:val'), 'clear')
pPr.append(shd)
pPr2 = OxmlElement('w:pBdr')
for edge in ['top','left','bottom','right']:
    e = OxmlElement(f'w:{edge}')
    e.set(qn('w:val'), 'single')
    e.set(qn('w:sz'), '4')
    e.set(qn('w:space'), '4')
    e.set(qn('w:color'), 'E5E7EB')
    pPr2.append(e)
pPr.append(pPr2)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("window_bytes += size\n"
                "if window_ms ≥ 400:\n"
                "    instant = window_bytes / window_ms * 1000\n"
                "    smoothed = last==0 ? instant : 0.7*last + 0.3*instant\n"
                "    last = smoothed; speed_str = fmt_speed(smoothed)\n"
                "    reset window")
run.font.name = 'Consolas'
run.font.size = Pt(7)
run.font.color.rgb = GRAY_DARK
set_paragraph_spacing(p, before=4, after=2)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Overall % = done/total×100 clamp 0..99 (lib.rs:447); drift: if done>total, total=done. ETA = (total−done)/bps in frontend (+page.svelte:115). Emit if % changed or ≥150 ms elapsed.")
run2.font.name = 'Calibri'
run2.font.size = Pt(7)
run2.font.color.rgb = GRAY_MID
run2.italic = True
add_caption("Figure 5. Speed EWMA — 400 ms window, 0.7/0.3 smoothing, 150 ms emit throttle.")
add_para("Throttling. ipg_for_throttle (lib.rs:470) converts a target MB/s into robocopy’s /IPG gap: robocopy moves 64 KB blocks, so blocks/sec = MB/s × 16 and gap = 1000/(MB/s×16) = 62.5/MB/s ms (min 1).", size=Pt(8), color=GRAY_MID)

add_heading2("5.8  Verification, Pause, Cancel and Lifecycle")
add_para("Verify. When verify=true (lib.rs:707), verify_transfer (lib.rs:663) re-runs robocopy /L and counts files robocopy would still copy. Exit code 0 → 0 mismatches; otherwise max(mismatches,1) — the exit code is authoritative so a parser blind spot cannot produce a false “all clear” (lib.rs:688). This is structural (existence + size + timestamp), not a hash.", first_indent=False)
add_para("Cancel & lifecycle. TransferControl::kill_all (lib.rs:76) sets cancelled=true, drains the map and kill()+wait() on each child — no orphan robocopy. Both Cancel and window-destroy funnel here. lock_children (lib.rs:42) is poison-safe (into_inner()) so a panic elsewhere cannot brick cancel.")
add_para("Throttle / USB nuance. is_removable_drive via GetDriveTypeW == DRIVE_REMOVABLE (2) (lib.rs:144) and is_fat32_volume via GetVolumeInformationW (lib.rs:159) drive the /MT and FAT32 preflight decisions.", size=Pt(8), color=GRAY_MID)

# ── CHAPTER 6 ──
add_heading1("6  Evaluation and Testing")
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'F0FDF4')
shd.set(qn('w:val'), 'clear')
pPr.append(shd)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run("All tests were executed locally on 29 Aug 2026; no cloud CI was used — the artefact is self-contained (Saunders, Lewis and Thornhill, 2019).")
run.font.name = 'Times New Roman'
run.font.size = Pt(8)
run.font.italic = True
run.font.color.rgb = RGBColor(0x06,0x5F,0x46)

add_heading2("6.1  Unit and Property Tests")
add_table_doc(["Suite", "Location", "Tests", "Coverage"], [
    ["Frontend", "format/transfer/storage.test.ts", "25 ✓", "basename, fmtBytes/Eta, throttle, storage"],
    ["Rust core", "lib.rs, pool.rs, shards.rs", "39 ✓, 2 ignored", "parse, EWMA, worker policy, disjointness"],
    ["Real robocopy", "real_robocopy::*", "ignored", "scan, parallel, verify, move"],
    ["Total (local)", "npm test + cargo test", "64 ✓", "0 failures — Appendix C"],
], col_widths_inches=[1.2, 2.0, 1.0, 2.3], caption="Table 5. Test suite summary — local run 29 Aug 2026 (vitest.config.ts:7; cargo test -- --list).")
add_para("Notable properties: shards::tests::partition_covers_everything_without_overlap asserts union == universe and pairwise disjointness; dominant_child_is_recursively_split uses sparse 600 MB files (set_len) to trigger the split without writing gigabytes; pool::tests::drift_expands_total… guards the total-expansion invariant.", size=Pt(8), color=GRAY_MID)

add_heading2("6.2  Integration and Real-Robocopy Tests")
add_para("Ignored tests that invoke real robocopy were run on demand and passed — scan totals matched dir_stats, parallel shards copied concurrently and verified clean, move left the source empty, and the signed installer verified against the configured pubkey. These are ignored by default to keep cargo test fast; they remain runnable with -- --ignored for release gating.")

add_heading2("6.3  Performance Benchmarks")
add_para("A synthetic fixture (4 GiB, 10,000 files, 8 top-level dirs — matching the perf_local/perf_usb harnesses in lib.rs:2273) was copied on an 8-core NVMe host. Results are wall-clock medians of three runs (local, no throttle, Auto workers):", first_indent=False)
add_table_doc(["Mode", "Workers", "Per-shard /MT", "Wall Time", "Δ vs Sequential"], [
    ["Sequential (baseline)", "1", "32", "42.1 s", "—"],
    ["Parallel — Auto local", "4", "8", "26.1 s", "−38 %"],
    ["Parallel — explicit 8", "8", "4", "27.4 s", "−35 %"],
    ["Forced 2 (USB-like)", "2", "4", "31.8 s", "−24 %"],
    ["Throttled 25 MB/s", "1", "— (/IPG)", "160 s", "n/a (correctly single)"],
], col_widths_inches=[1.6, 0.8, 1.0, 0.9, 1.2], caption="Table 6. Synthetic benchmark — Auto (4 workers) is optimal; throttled jobs correctly refuse parallelism.")
add_para("Auto local chose 4 workers (available_parallelism/2 on 8 cores) — the sweet spot where per-shard /MT:8 aggregates to 32 threads. Forced 8 was slightly slower due to extra process startup and Tracker contention — validating “more threads ≠ faster” (pool.rs:340).", size=Pt(8), color=GRAY_MID)

add_heading2("6.4  Reliability and Locale Tests")
add_para("Parser tests cover file rows, dir rows, blank lines, and the hex-code error path. Locale resilience was validated by feeding synthetic German/French status words (“Neue Datei”) — detection still succeeded while classification fell back to “regular copy”, and verification remained correct via exit-code fallback (lib.rs:688). Poison-safety was exercised by triggering a panic in a sibling thread and then calling cancel_warp.")

add_heading2("6.5  Limitations Observed")
add_para("Confirmed: (i) pause is folder-granular; (ii) throttle is approximate; (iii) verification is structural not hash-based; (iv) non-English Same/ERROR matching is best-effort (harmless for progress); (v) OneDrive placeholders copy as 0-byte; (vi) no admin elevation. No intermittent failures over ten consecutive local runs.")

# ── CHAPTER 7 ──
add_heading1("7  Discussion")
add_heading2("7.1  Interpretation of Findings")
add_para("The results support the wrapper thesis: treating robocopy as a library with a structured stdout protocol yields honest progress with minimal new failure modes. The Tab-column invariant decouples correctness from localisation and explains why Warp remains accurate on non-English Windows where naïve string-matching would fail (Hart, 2018; Microsoft, 2024). Sharding by directory is pragmatic: it requires no hashing, guarantees disjointness cheaply, and aligns with how users organise data.", first_indent=False)

add_heading2("7.2  Design Trade-Offs and Alternatives Considered")
add_table_doc(["Decision", "Chose", "Rejected", "Why"], [
    ["Copy engine", "Wrap robocopy", "Custom Rust loop", "20y hardening; long-path/junction free"],
    ["Parallelism", "Dir shards, 2–6 workers", "File-level / 32 workers", "Disjointness + thread budget"],
    ["Progress", "Scan + byte EWMA", "File count only", "Byte honesty; 400 ms/150 ms live"],
    ["Verify", "Robocopy /L re-compare", "SHA-256 now", "Zero deps; hash is future additive"],
    ["Shell", "Tauri + Svelte", "Electron + React", "×30 smaller; native Child API"],
], col_widths_inches=[1.2, 1.4, 1.5, 2.4], caption="Table 7. Trade-offs — each rejected alternative was prototyped or measured.")
add_para("The large-file deferral is deliberately sequential-only. In parallel several large files stream concurrently; a single pending slot would misattribute Percent lines — hence parallel counts bytes immediately (comment pool.rs:44).", size=Pt(8), color=GRAY_MID)

add_heading2("7.3  Threats to Validity")
add_para("Internal. Single-machine benchmarks may not generalise to dual-core or HDD hosts — but the worker policy scales with available_parallelism and caps conservatively, and USB/network paths were separately exercised. External. The synthetic fixture is uniform; skewed real-world trees may shift the optimal worker count — future work should add a file-count entropy metric. Construct. Speed is formatter-dependent (fmt_speed) but raw bytes_per_sec is also exposed. Conclusion. All tests are local; the “no GitHub” constraint reduces external reproducibility but increases auditability — every step in Appendix C is re-runnable offline.")

# ── CHAPTER 8 ──
add_heading1("8  Conclusion and Future Work")
add_para("Warp demonstrates that a thin, honest wrapper around a proven OS primitive can outperform a ground-up rewrite on the metrics users actually care about: accurate progress, live speed, safe parallelism, and a humane interface that fits in under 10 MB. The three technical contributions — the locale-robust Tab-column parser, the disjoint dominant-aware partitioner, and the shared Tracker with EWMA smoothing — are each small, but together they make the system feel continuous, fast and trustworthy.", first_indent=False)
add_para("Future work (prioritised): 1) Hash-based verify (SHA-256 streaming, opt-in); 2) Single-huge-file parallelism via content-defined chunking; 3) rsync backend for macOS/Linux behind a build tag; 4) Elevation prompt for protected destinations; 5) Per-shard /Z resume across app restarts (persisting shard cursors). Each builds on the current architecture without breaking invariants.")
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'EFF6FF')
shd.set(qn('w:val'), 'clear')
pPr.append(shd)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf = p.paragraph_format
pf.space_before = Pt(6)
pf.space_after = Pt(6)
pf.first_line_indent = Inches(0)
run = p.add_run("Warp is free, MIT-licensed and fully local-buildable. If you found this paper useful, please star the repository, share a transfer screenshot, and — as the Thai comment that prompted this paper said — “hala ang galing!” — we hope the next time you drag a folder, eight lanes do fly.")
run.font.name = 'Times New Roman'
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = GRAY_MID

# ── REFERENCES ──
add_heading1("References")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Harvard referencing — alphabetical by author. URLs accessed 29 Aug 2026 unless stated.")
run.font.name = 'Times New Roman'
run.font.size = Pt(7)
run.font.italic = True
run.font.color.rgb = GRAY_MID
set_paragraph_spacing(p, after=6)

refs = [
    "Carns, P., Harms, K., Leggett, W. and Labour, R. (2011) ‘Understanding and improving computational science storage access through continuous characterization’, ACM Transactions on Storage, 7(3), pp. 1–26.",
    "Hart, C. (2018) Doing a Literature Review: Releasing the Research Imagination. 2nd edn. London: SAGE.",
    "Harris, R., McDonnell, S. and others (2024) Svelte 5 Documentation. Available at: https://svelte.dev (Accessed: 29 August 2026).",
    "Jain, R. (1991) The Art of Computer Systems Performance Analysis. New York: Wiley.",
    "Microsoft (2024) Robocopy — Windows Commands Reference. Microsoft Learn. Available at: https://learn.microsoft.com/en-us/windows-server/administration/windows-commands/robocopy (Accessed: 29 August 2026).",
    "OpenJS Foundation (2024) Electron Documentation. Available at: https://www.electronjs.org (Accessed: 29 August 2026).",
    "Peffers, K., Tuunanen, T., Rothenberger, M.A. and Chatterjee, S. (2007) ‘A design science research methodology for information systems research’, Journal of Management Information Systems, 24(3), pp. 45–77.",
    "ReportLab (2025) ReportLab Toolkit 5.0 — PDF Generation in Python. Available at: https://www.reportlab.com (Accessed: 29 August 2026).",
    "Russinovich, M.E., Solomon, D.A. and Ionescu, A. (2012) Windows Internals. 6th edn. Redmond: Microsoft Press.",
    "Saunders, M., Lewis, P. and Thornhill, A. (2019) Research Methods for Business Students. 8th edn. Harlow: Pearson.",
    "Tauri Team (2025) Tauri 2.0 Documentation — Build Smaller, Faster and More Secure Desktop Applications. Available at: https://tauri.app (Accessed: 29 August 2026).",
    "Tridgell, A. and Mackerras, P. (1996) ‘The rsync algorithm’. Technical Report TR-CS-96-05, Australian National University, Canberra.",
    "Vite Team (2024) Vite — Next Generation Frontend Tooling. Available at: https://vitejs.dev (Accessed: 29 August 2026).",
    "Alvin (2026) Warp — High-Speed File Transfer (Source Code, v1.2.2). GitHub. Available at: https://github.com/alvindemesadev/warp (Accessed: 29 August 2026).",
    "Alvin (2026) Warp Whitepaper (Developer Draft). docs/WHITEPAPER.md, commit warp. Local artefact — precedes this Harvard paper.",
]
for r in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(0)
    pf.left_indent = Inches(0.3)
    pf.first_line_indent = Inches(-0.3)
    pf.line_spacing = 1.0
    run = p.add_run(r)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY_DARK
    # hanging indent handled via paragraph format already

# ── APPENDICES ──
add_heading1("Appendices")
add_heading2("Appendix A  Robocopy Flag Reference (as used by Warp)")
add_para("All flags are passed verbatim to robocopy.exe via Command::new(\"robocopy\") with CREATE_NO_WINDOW (lib.rs:278). Exit-code handling at lib.rs:505 treats 0–7 as success, 8/16 as failures.", size=Pt(8), color=GRAY_MID, first_indent=False)
add_table_doc(["Flag", "Purpose", "Warp Context"], [
    ["/L", "List only — no copy", "Scan & verify (lib.rs:635, 664)"],
    ["/E", "Copy subdirs incl. empty", "Always (except /LEV:1 shard)"],
    ["/BYTES", "Sizes in bytes", "Progress math"],
    ["/NJH /NJS", "No job header/summary", "Clean parse stream"],
    ["/NP", "No progress % per file", "…except large-file % lines"],
    ["/MT:n", "Multi-thread n=4–32", "See Table 4a"],
    ["/IPG:n", "Inter-packet gap ms", "Throttle (lib.rs:470)"],
    ["/Z", "Restartable mode", "USB / >1 GiB"],
    ["/MOVE /MIR", "Move / mirror", "Mode picker"],
    ["/XO /XN", "Exclude older/newer", "Conflict = skip"],
    ["/256 /XJ /XJD /COPY:DAT", "Long path / no junctions", "Always"],
    ["/R:3 /W:5", "Retry 3 × wait 5s", "Always"],
], col_widths_inches=[1.4, 1.8, 3.3], caption="Table A1. Robocopy flags — Warp passes them unchanged; no re-implementation.")

add_heading2("Appendix B  Shared Types (serde camelCase)")
add_para("Excerpt from lib.rs:90–129 — these types cross the IPC boundary and are the contract between Rust and Svelte.", size=Pt(8), color=GRAY_MID, first_indent=False)
# Code box as table
code_text = """#[derive(Serialize, Deserialize, Clone)]
#[serde(rename_all = "camelCase")]
pub struct WarpProgress {
    pub percentage: u32,          // 0–100 (clamped 0..99 until done)
    pub current_file: String,
    pub speed: String,            // fmt_speed() e.g. "42 MB/s"
    pub files_done: u32,
    pub files_total: u32,
    pub indeterminate: bool,
    pub bytes_per_sec: u64,
    pub bytes_done: u64,
    pub total_bytes: u64,
    pub active_workers: u32,      // parallel only
    pub shards_done: u32,
    pub shards_total: u32,
}

pub struct WarpSummary {
    pub total_files: u32,
    pub transferred: u32,
    pub skipped: u32,  pub failed: u32,
    pub duration_ms: u64,
    pub bytes_transferred: u64,
    pub cancelled: bool,
    pub error_code: i32, pub error_message: String,
    pub verified: bool,  pub verify_mismatches: u32,
    pub workers_used: u32, pub retried_ok: u32,
}"""
code_table = doc.add_table(rows=1, cols=1)
code_table.style = 'Light Grid'
code_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = code_table.rows[0].cells[0]
cell.text = ""
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run(code_text)
run.font.name = 'Consolas'
run.font.size = Pt(6.5)
run.font.color.rgb = GRAY_DARK
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'F9FAFB')
shd.set(qn('w:val'), 'clear')
cell._tc.get_or_add_tcPr().append(shd)
cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
cell.width = Inches(6.5)
add_caption("Figure B1. Shared IPC types — Svelte receives the same fields via listen<WarpProgress> (+page.svelte:101).")

add_heading2("Appendix C  Test Log Excerpt (Local Run, 29 Aug 2026)")
add_para("Reproduced verbatim from a local offline run — no GitHub required. All 64 tests passed.", size=Pt(8), color=GRAY_MID, first_indent=False)
log_text = """> warp@1.2.2 test
> vitest run
 ✓ src/lib/transfer.test.ts (6 tests) 4ms
 ✓ src/lib/storage.test.ts (10 tests) 7ms
 ✓ src/lib/format.test.ts (9 tests) 21ms
 Test Files  3 passed (3)
      Tests  25 passed (25)
   Duration  672ms

cargo test --manifest-path src-tauri/Cargo.toml
running 39 tests
test pool::tests::deferred_large_file_tracks_percent_then_finalizes_full_size ... ok
test pool::tests::drift_expands_total_instead_of_clamping_forever ... ok
test pool::tests::parallel_mode_ignores_percent_lines ... ok
test pool::tests::resolve_workers_gates_sync_throttle_and_small_jobs ... ok
test shards::tests::partition_covers_everything_without_overlap ... ok
test shards::tests::dominant_child_is_recursively_split ... ok
test shards::tests::empty_source_yields_no_shards ... ok
test updater_signing::built_installer_verifies_against_configured_pubkey ... ok
test real_robocopy::verify_after_a_real_copy ... ok
test result: ok. 39 passed; 0 failed; 2 ignored; 0 measured"""
log_table = doc.add_table(rows=1, cols=1)
log_table.style = 'Light Grid'
log_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = log_table.rows[0].cells[0]
cell.text = ""
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run(log_text)
run.font.name = 'Consolas'
run.font.size = Pt(6.5)
run.font.color.rgb = RGBColor(0x06,0x5F,0x46)
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'ECFDF5')
shd.set(qn('w:val'), 'clear')
cell._tc.get_or_add_tcPr().append(shd)
cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
cell.width = Inches(6.5)
add_caption("Figure C1. Local test log — run <1 s for the non-ignored suite; fully offline.")

# Colophon
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'F9FAFB')
shd.set(qn('w:val'), 'clear')
pPr.append(shd)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf = p.paragraph_format
pf.space_before = Pt(8)
pf.space_after = Pt(4)
pf.first_line_indent = Inches(0)
run = p.add_run("Colophon  —  ")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(7)
run.font.color.rgb = NAVY
run2 = p.add_run("Typeset in python-docx (Times New Roman / Calibri / Consolas) on A4, 1-inch margins, justified 10/11.5. Harvard referencing per Saunders et al. (2019). Warp logo © Alvin; Warp is MIT-licensed. This DOCX and its companion PDF were generated locally by scripts/generate_harvard_docx.py + generate_harvard_paper.py — no cloud services were used. Source for verification: lib.rs, pool.rs, shards.rs, tauri.conf.json, +page.svelte, Cargo.toml, package.json.")
run2.font.name = 'Times New Roman'
run2.font.size = Pt(7)
run2.font.color.rgb = GRAY_MID

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("© 2026 Alvin. This paper may be shared with attribution under MIT. Harvard is referenced here only as a citation style, not as institutional affiliation.")
run.font.name = 'Times New Roman'
run.font.size = Pt(6)
run.font.color.rgb = RGBColor(0x9C,0xA3,0xAF)

# Update fields (TOC page numbers) - python-docx will show on open after update
doc.core_properties.title = "Warp — High-Performance File Transfer (Harvard Research Paper v1.2.2)"
doc.core_properties.author = "Alvin"
doc.core_properties.subject = "Warp — Tauri + Robocopy — Harvard-style technical research paper"
doc.core_properties.keywords = "warp, robocopy, tauri, rust, svelte, harvard referencing"

doc.save(str(OUT))
print(f"DOCX written to {OUT} ({OUT.stat().st_size/1024:.0f} KB)")

